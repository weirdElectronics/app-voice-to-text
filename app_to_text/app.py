from flask import Flask, request, render_template
import speech_recognition as sr
from docx import Document
import os
import base64
import io
import openpyxl
import re

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/guardar_audio', methods=['POST'])
def guardar_audio():
    audio_b64 = request.form['audio']
    modo = request.form.get('modo', 'texto')
    audio_data = audio_b64.split(',')[1]
    audio_bytes = base64.b64decode(audio_data)

    webm_path = "grabacion.webm"
    wav_path = "grabacion.wav"
    with open(webm_path, "wb") as f:
        f.write(audio_bytes)

    os.system(f"ffmpeg -y -i {webm_path} -ar 16000 -ac 1 -f wav {wav_path}")

    r = sr.Recognizer()
    with sr.AudioFile(wav_path) as source:
        audio_rec = r.record(source)
    try:
        texto = r.recognize_google(audio_rec, language="es-AR")
    except Exception as e:
        return f"Error al transcribir: {e}"

    if modo == "texto":
        doc_path = os.path.expanduser("~/Desktop/transcripciones.docx")
        if os.path.exists(doc_path):
            doc = Document(doc_path)
        else:
            doc = Document()
            doc.add_heading("Transcripciones de voz", 0)

        doc.add_paragraph(texto)
        doc.save(doc_path)
        return f"Texto guardado en documento: {texto}"

    elif modo == "suma":
        excel_path = os.path.expanduser("~/Desktop/gastos.xlsx")
        if os.path.exists(excel_path):
            wb = openpyxl.load_workbook(excel_path)
            ws = wb.active
        else:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Gastos"
            ws.append(["Descripción", "Monto"])

        match = re.search(r"(\d+(?:[.,]\d+)*)", texto)
        if match:
            monto_str = match.group(1).replace(",", ".")
            try:
                monto = float(monto_str)
            except:
                monto = 0.0
        else:
            monto = 0.0

        descripcion = texto.replace(match.group(1), "").strip() if match else texto

        ws.append([descripcion, monto])

        # Actualizar total en una fila aparte
        total_row = ws.max_row + 1
        ws[f"A{total_row}"] = "TOTAL"
        ws[f"B{total_row}"] = f"=SUM(B2:B{ws.max_row})"

        wb.save(excel_path)
        return f"Gasto registrado: {descripcion} (monto: {monto})"

@app.route('/reset_documento', methods=['POST'])
def reset_documento():
    doc_path = os.path.expanduser("~/Desktop/transcripciones.docx")
    if os.path.exists(doc_path):
        os.remove(doc_path)

    doc = Document()
    doc.add_heading("Transcripciones de voz", 0)
    doc.save(doc_path)

    excel_path = os.path.expanduser("~/Desktop/gastos.xlsx")
    if os.path.exists(excel_path):
        os.remove(excel_path)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Gastos"
    ws.append(["Descripción", "Monto"])
    wb.save(excel_path)

    return "Documentos reiniciados. Word y Excel están vacíos y listos para nuevas transcripciones."

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
