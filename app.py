from flask import Flask, request, render_template, send_file, session
import speech_recognition as sr
from docx import Document
import os
import base64
import openpyxl
import re
import uuid

app = Flask(__name__, static_folder="static", template_folder="templates")
app.secret_key = "clave_supersecreta"  # Necesario para manejar sesiones

DATA_DIR = "./data"
os.makedirs(DATA_DIR, exist_ok=True)

# Función para obtener rutas de archivos según el usuario
def get_paths():
    user_id = session.get("user_id")
    if not user_id:
        user_id = str(uuid.uuid4())
        session["user_id"] = user_id
    word_path = os.path.join(DATA_DIR, f"transcripciones_{user_id}.docx")
    excel_path = os.path.join(DATA_DIR, f"gastos_{user_id}.xlsx")
    return word_path, excel_path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/guardar_audio', methods=['POST'])
def guardar_audio():
    audio_b64 = request.form['audio']
    modo = request.form.get('modo', 'texto')
    audio_data = audio_b64.split(',')[1]
    audio_bytes = base64.b64decode(audio_data)

    webm_path = os.path.join(DATA_DIR, "grabacion.webm")
    wav_path = os.path.join(DATA_DIR, "grabacion.wav")
    with open(webm_path, "wb") as f:
        f.write(audio_bytes)

    os.system(f"ffmpeg -y -i {webm_path} -ar 16000 -ac 1 -f wav {wav_path}")

    r = sr.Recognizer()
    with sr.AudioFile(wav_path) as source:
        audio_rec = r.record(source)
    try:
        texto = r.recognize_google(audio_rec, language="es-AR")
    except Exception as e:
        return f"Error al transcribir: {e}"

    WORD_PATH, EXCEL_PATH = get_paths()

    if modo == "texto":
        if os.path.exists(WORD_PATH):
            doc = Document(WORD_PATH)
        else:
            doc = Document()

        p = doc.add_paragraph(texto)
        run = p.runs[0]
        run.font.name = "Courier New"

        doc.save(WORD_PATH)
        return f"Texto guardado en documento: {texto}"

    elif modo == "suma":
        if os.path.exists(EXCEL_PATH):
            wb = openpyxl.load_workbook(EXCEL_PATH)
            ws = wb.active
        else:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Gastos"
            ws.append(["Descripción", "Monto"])

        match = re.search(r"(\d+(?:[.,]\d+)*)", texto.lower())
        if match:
            monto_str = match.group(1).replace(",", ".")
            try:
                monto = float(monto_str)
            except:
                monto = 0.0
            if "mil" in texto.lower():
                monto *= 1000
        else:
            monto = 0.0

        descripcion = texto.replace(match.group(1), "").strip() if match else texto
        ws.append([descripcion, monto])

        total = sum(
            cell.value for cell in ws["B"][1:] if isinstance(cell.value, (int, float))
        )
        ws["A1"] = "TOTAL"
        ws["B1"] = total

        wb.save(EXCEL_PATH)
        return f"Gasto registrado: {descripcion} (monto: {monto})"

@app.route('/reset_documento', methods=['POST'])
def reset_documento():
    WORD_PATH, EXCEL_PATH = get_paths()

    if os.path.exists(WORD_PATH):
        os.remove(WORD_PATH)
    doc = Document()
    doc.save(WORD_PATH)

    if os.path.exists(EXCEL_PATH):
        os.remove(EXCEL_PATH)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Gastos"
    ws.append(["Descripción", "Monto"])
    wb.save(EXCEL_PATH)

    return "Tus documentos fueron reiniciados. Word y Excel están vacíos y listos para nuevas transcripciones."

@app.route('/descargar_word')
def descargar_word():
    WORD_PATH, _ = get_paths()
    if os.path.exists(WORD_PATH):
        return send_file(WORD_PATH,
                         as_attachment=True,
                         download_name="transcripciones.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    return "No hay documento Word disponible."

@app.route('/descargar_excel')
def descargar_excel():
    _, EXCEL_PATH = get_paths()
    if os.path.exists(EXCEL_PATH):
        return send_file(EXCEL_PATH,
                         as_attachment=True,
                         download_name="gastos.xlsx",
                         mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    return "No hay documento Excel disponible."

# NUEVAS RUTAS PARA VER ONLINE
@app.route('/ver_word')
def ver_word():
    WORD_PATH, _ = get_paths()
    if os.path.exists(WORD_PATH):
        doc = Document(WORD_PATH)
        contenido = [p.text for p in doc.paragraphs if p.text.strip()]
        return render_template("ver_word.html", contenido=contenido)
    return "No hay documento Word disponible."

@app.route('/ver_excel')
def ver_excel():
    _, EXCEL_PATH = get_paths()
    if os.path.exists(EXCEL_PATH):
        wb = openpyxl.load_workbook(EXCEL_PATH)
        ws = wb.active
        filas = [[cell.value for cell in row] for row in ws.iter_rows()]
        return render_template("ver_excel.html", filas=filas)
    return "No hay documento Excel disponible."

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)


